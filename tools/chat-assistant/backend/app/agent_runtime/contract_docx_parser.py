"""Native DOCX parser for contract documents."""

from __future__ import annotations

from pathlib import Path

from app.services.document_parser_types import ParsedBlock


def parse_docx_blocks(file_path: str) -> list[ParsedBlock]:
    """Extract paragraphs and tables from a DOCX file.

    DOCX native text is treated as the primary source for contract clauses.
    Layout/page evidence can be added later by the MinerU validation stage.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX 解析需要安装 python-docx，请执行 pip install -r requirements.txt") from exc

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {file_path}")

    document = Document(str(path))
    blocks: list[ParsedBlock] = []
    current_title = ""

    for paragraph in document.paragraphs:
        text = _normalize(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if _looks_like_heading(style_name, text):
            current_title = text[:120]
        blocks.append(ParsedBlock(text=text, section_title=current_title))

    for table_index, table in enumerate(document.tables, 1):
        rows = []
        for row in table.rows:
            cells = [_normalize(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(ParsedBlock(
                text="\n".join(rows),
                section_title=current_title or f"表格 {table_index}",
            ))

    return blocks


def _normalize(text: str) -> str:
    lines = [line.strip() for line in str(text or "").replace("\r", "\n").splitlines()]
    return "\n".join(line for line in lines if line)


def _looks_like_heading(style_name: str, text: str) -> bool:
    style = (style_name or "").lower()
    if "heading" in style or "标题" in style:
        return True
    if len(text) <= 80 and (
        text.startswith("第") and "条" in text[:12]
        or text[:2].isdigit() and any(text.startswith(prefix) for prefix in ("1.", "2.", "3.", "4.", "5."))
    ):
        return True
    return False
